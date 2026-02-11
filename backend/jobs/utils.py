import re
import io
from pdfminer.high_level import extract_text
import docx
import os

# A broad list of common professional and technical skills
# Professional and Technical Skills categorized
TECHNICAL_SKILLS = {
    'Python', 'JavaScript', 'Java', 'C++', 'C#', 'PHP', 'Ruby', 'Swift', 'Kotlin', 'Go', 'Rust', 'TypeScript', 'SQL', 'HTML', 'CSS',
    'Django', 'React', 'Angular', 'Vue', 'Spring', 'Flask', 'Laravel', 'Express', 'Bootstrap', 'jQuery', 'Node.js', 'TensorFlow', 'PyTorch', 'FastAPI',
    'AWS', 'Azure', 'Google Cloud', 'Docker', 'Kubernetes', 'CI/CD', 'Jenkins', 'Terraform', 'PostgreSQL', 'MySQL', 'MongoDB', 'Redis', 'Nginx',
    'Git', 'GitHub'
}

SOFT_SKILLS = {
    'Project Management', 'UI/UX Design', 'SEO', 'Data Analysis', 'Machine Learning', 'AI', 'Graphic Design', 'Excel', 'Agile', 'Scrum',
    'Marketing', 'Sales', 'Communication', 'Leadership', 'Content Writing', 'Research', 'Public Speaking', 'Problem Solving', 'Teamwork',
    'Time Management', 'Critical Thinking', 'Adaptability', 'Creativity'
}

LANGUAGES = {
    'English', 'Spanish', 'French', 'German', 'Chinese', 'Japanese', 'Arabic', 'Russian', 'Portuguese', 'Italian', 'Albanian'
}

SKILLS_LIST = TECHNICAL_SKILLS.union(SOFT_SKILLS).union(LANGUAGES)

def extract_text_from_file(file_obj):
    """Extracts text from PDF or DOCX file object by trying multiple libraries."""
    file_extension = file_obj.name.split('.')[-1].lower()
    text = ""
    
    try:
        # Read file content into BytesIO to ensure it's a standard file-like object
        import io
        file_obj.seek(0)
        content = file_obj.read()
        f_io = io.BytesIO(content)
        
        if file_extension == 'pdf':
            # Try pdfminer.six first (best for layout)
            try:
                from pdfminer.high_level import extract_text
                f_io.seek(0)
                text = extract_text(f_io)
            except Exception as e:
                print(f"pdfminer.six failed: {e}")
            
            # Fallback to pypdf if pdfminer returns nothing meaningful
            if not text or not text.strip() or len(text.strip()) < 20:
                try:
                    from pypdf import PdfReader
                    f_io.seek(0)
                    reader = PdfReader(f_io)
                    pypdf_text = ""
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            pypdf_text += page_text + "\n"
                    if pypdf_text.strip():
                        print("Fallback to pypdf successful.")
                        text = pypdf_text
                except Exception as e:
                    print(f"pypdf fallback failed: {e}")

        elif file_extension == 'docx':
            f_io.seek(0)
            doc = docx.Document(f_io)
            
            # Extract text from paragraphs
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Extract text from tables (Crucial for CVs often designed with tables)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" | ".join(row_text))
            
            text = "\n".join(full_text)
        
        # Clean up text: remove null bytes and strip whitespace
        if text:
            text = text.replace('\x00', '').strip()
            
    except Exception as e:
        import traceback
        print(f"Error extracting text from {file_obj.name}: {str(e)}")
        traceback.print_exc()
        
    return text

def parse_cv_data(text):
    """
    Parses text to find phone, location, linkedin, and github.
    Very basic regex-based extraction.
    """
    data = {
        'first_name': None,
        'last_name': None,
        'phone': None,
        'location': None,
        'linkedin': None,
        'github': None,
        'summary': None,
        'technical_skills': [],
        'soft_skills': [],
        'languages': [],
    }
    
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    if lines:
        # Heuristic: the first line is often the name
        name_line = lines[0]
        name_parts = name_line.split()
        if len(name_parts) >= 2:
            data['first_name'] = name_parts[0]
            data['last_name'] = " ".join(name_parts[1:])
        elif len(name_parts) == 1:
            data['first_name'] = name_parts[0]

    # Phone number regex (flexible)
    phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
    phone_match = re.search(phone_pattern, text)
    if phone_match:
        data['phone'] = phone_match.group()
        
    # LinkedIn
    linkedin_pattern = r'linkedin\.com/in/[\w-]+'
    linkedin_match = re.search(linkedin_pattern, text)
    if linkedin_match:
        data['linkedin'] = "https://www." + linkedin_match.group()
        
    # GitHub
    github_pattern = r'github\.com/[\w-]+'
    github_match = re.search(github_pattern, text)
    if github_match:
        # Avoid matching generic github links if possible, but for a CV it's usually the profile
        data['github'] = "https://www." + github_match.group()

    # Location (looking for common patterns like "City, Country" or "City, State")
    lines = text.split('\n')
    for line in lines:
        stripped_line = line.strip()
        if not stripped_line:
            continue
        if any(keyword in stripped_line.lower() for keyword in ['phone', 'email', 'linkedin', 'github']):
            continue
        if ',' in stripped_line and len(stripped_line) < 50:
            data['location'] = stripped_line
            break

    # Summary
    summary_match = re.search(r'(?:Summary|Profile|About Me|Professional Summary)[:\s]*(.*?)(?=\n\s*(?:Work Experience|Experience|Employment|Education|Skills|Projects|Languages|Certifications)|$)', text, re.IGNORECASE | re.DOTALL)
    if summary_match:
        data['summary'] = summary_match.group(1).strip()
    else:
        # Fallback for summaries that are just text at the top
        summary_lines = []
        started = False
        for line in lines:
            if any(h in line.lower() for h in ['summary', 'profile', 'about me']):
                started = True
                content = re.sub(r'^(?:summary|profile|about me|professional summary)[:\s]*', '', line, flags=re.IGNORECASE).strip()
                if content: summary_lines.append(content)
                continue
            if started:
                if any(h in line.lower() for h in ['experience', 'education', 'skills', 'projects']):
                    break
                summary_lines.append(line)
        if summary_lines:
            data['summary'] = " ".join(summary_lines).strip()
    
    # Experience Extraction
    data['experiences'] = []
    experience_section = re.search(r'(?:Work Experience|Experience|Employment|Work History)[:\s]*(.*?)(?=\n\s*(?:Education|Skills|Projects|Languages|Certifications|Interests)|$)', text, re.IGNORECASE | re.DOTALL)
    
    if experience_section:
        exp_text = experience_section.group(1).strip()
        
        # Heuristic: Experiences often start with a line containing a date or follow a pattern
        # Let's split by lines that look like they start with a date range
        # e.g. "Jan 2020 - " or "2020 - 2022"
        entry_starts = []
        lines = exp_text.split('\n')
        for i, line in enumerate(lines):
            line = line.strip()
            # Look for date patterns at the beginning or end of the line
            if re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2}).*?(?:Present|\d{4})', line, re.IGNORECASE):
                entry_starts.append(i)
        
        # If we didn't find clear date starts, just treat the whole block as one entry
        if not entry_starts:
            entry_starts = [0]
            
        for i in range(len(entry_starts)):
            start_idx = entry_starts[i]
            end_idx = entry_starts[i+1] if i+1 < len(entry_starts) else len(lines)
            
            entry_lines = [l.strip() for l in lines[start_idx:end_idx] if l.strip()]
            if not entry_lines: continue
            
            current_exp = {
                'company': 'Unknown Company',
                'position': 'Unknown Position',
                'date_str': '',
                'description': ''
            }
            
            # Find the date line in this entry
            date_line_idx = -1
            for j, l in enumerate(entry_lines):
                if re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2}).*?(?:Present|\d{4})', l, re.IGNORECASE):
                    current_exp['date_str'] = l
                    date_line_idx = j
                    break
            
            # Heuristic for position and company:
            # Usually the line before or the line with the date, or the first line
            other_lines = [l for j, l in enumerate(entry_lines) if j != date_line_idx]
            
            if other_lines:
                first_line = other_lines[0]
                if ' at ' in first_line:
                    parts = first_line.split(' at ')
                    current_exp['position'] = parts[0].strip()
                    current_exp['company'] = parts[1].strip()
                elif ' | ' in first_line:
                    parts = first_line.split(' | ')
                    current_exp['position'] = parts[0].strip()
                    current_exp['company'] = parts[1].strip()
                else:
                    current_exp['position'] = first_line
                    if len(other_lines) > 1:
                        current_exp['company'] = other_lines[1]
                
                current_exp['description'] = "\n".join(other_lines[1:]) if len(other_lines) > 1 else ""
            
            data['experiences'].append(current_exp)

    # Extract Skills categories
    data['technical_skills'] = extract_skills(text, TECHNICAL_SKILLS)
    data['soft_skills'] = extract_skills(text, SOFT_SKILLS)
    data['languages'] = extract_skills(text, LANGUAGES)

    return data

def extract_skills(text, skills_list=None):
    """Matches text against a provided list of skills. If skills_list is None, uses TECHNICAL_SKILLS."""
    if not text:
        return []
    
    if skills_list is None:
        skills_list = TECHNICAL_SKILLS
        
    found_skills = []
    for skill in skills_list:
        pattern = r'\b' + re.escape(skill) + r'\b'
        if re.search(pattern, text, re.IGNORECASE):
            found_skills.append(skill)
            
    return sorted(list(set(found_skills)))

def calculate_match_score(user_skills_str, job_skills_str):
    """
    Calculates a match score between 0 and 100 based on overlapping skills.
    Skills are stored as comma-separated strings.
    """
    if not user_skills_str or not job_skills_str:
        return 0
        
    user_skills = set(s.strip().lower() for s in user_skills_str.split(',') if s.strip())
    job_skills = set(s.strip().lower() for s in job_skills_str.split(',') if s.strip())
    
    if not job_skills:
        return 0
        
    matches = user_skills.intersection(job_skills)
    score = (len(matches) / len(job_skills)) * 100
    
    return round(score)

def process_cv_and_update_profile(user, cv_file):
    """
    Centralized function to process a CV file, update the CV model,
    and auto-fill the UserProfile.
    """
    from .models import CV, UserProfile
    
    # Save or update CV model
    cv, created = CV.objects.get_or_create(user=user)
    cv.resume_file = cv_file
    if not cv.title:
        cv.title = f"CV - {user.username}"
    cv.save()

    # Extract text and parse data
    try:
        text = extract_text_from_file(cv_file)
        if text:
            parsed_data = parse_cv_data(text)
            
            # Update UserProfile
            profile, created = UserProfile.objects.get_or_create(user=user)
            
            # Sync User fields with extracted data
            if parsed_data.get('first_name'):
                user.first_name = parsed_data['first_name']
            if parsed_data.get('last_name'):
                user.last_name = parsed_data['last_name']
            if parsed_data.get('first_name') or parsed_data.get('last_name'):
                user.save()

            # Sync profile fields with extracted data
            if parsed_data.get('phone'):
                profile.phone = parsed_data['phone']
            if parsed_data.get('location'):
                profile.location = parsed_data['location']
            profile.save()
            
            # Update CV model specialized fields - update if newer or force update
            if parsed_data.get('summary'):
                cv.summary = parsed_data['summary']
            if parsed_data.get('linkedin'):
                cv.linkedin = parsed_data['linkedin']
            if parsed_data.get('github'):
                cv.github = parsed_data['github']
            
            # Update Skill categories - Re-calculate to avoid miscategorization carryover
            if parsed_data.get('technical_skills'):
                cv.skills = ", ".join(parsed_data['technical_skills'])
            else:
                cv.skills = ""
                
            if parsed_data.get('soft_skills'):
                cv.soft_skills = ", ".join(parsed_data['soft_skills'])
            else:
                cv.soft_skills = ""

            if parsed_data.get('languages'):
                cv.languages = ", ".join(parsed_data['languages'])
            else:
                cv.languages = ""
            
            cv.save()

            # Handle Experiences - Clear and re-populate
            if parsed_data.get('experiences'):
                from .models import Experience
                from datetime import datetime
                cv.experiences.all().delete()
                
                for exp in parsed_data['experiences']:
                    # Enhanced date parsing heuristic
                    start_date = None
                    end_date = None
                    date_str = exp.get('date_str', '')
                    
                    def parse_flexible_date(s):
                        if not s: return None
                        # Try month/year
                        for fmt in ('%b %Y', '%B %Y', '%m/%Y', '%Y'):
                            try:
                                # Clean string from common artifacts
                                clean_s = re.search(r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|20\d{2}).*?(?:\d{4}|Present)', s, re.IGNORECASE)
                                if clean_s:
                                    return datetime.strptime(clean_s.group(), fmt).date()
                            except:
                                continue
                        return None

                    # Split date_str by common separators
                    parts = re.split(r'[-–—]| to ', date_str, flags=re.IGNORECASE)
                    if len(parts) >= 1:
                        start_date = parse_flexible_date(parts[0].strip())
                    if len(parts) >= 2:
                        end_part = parts[1].strip().lower()
                        if 'present' in end_part or 'now' in end_part:
                            end_date = None
                        else:
                            end_date = parse_flexible_date(end_part)
                    
                    # Last resort fallback to years
                    if not start_date:
                        years = re.findall(r'20\d{2}', date_str)
                        if years:
                            try:
                                start_date = datetime.strptime(years[0], '%Y').date()
                                if len(years) >= 2:
                                    end_date = datetime.strptime(years[1], '%Y').date()
                            except: pass
                    
                    # Final safety check
                    if not start_date:
                        start_date = datetime.now().date()

                    Experience.objects.create(
                        cv=cv,
                        company=exp.get('company', 'Unknown'),
                        position=exp.get('position', 'Staff'),
                        start_date=start_date,
                        end_date=end_date,
                        description=exp.get('description', '')
                    )
            
            return True, "CV uploaded and profile updated with extracted information!"
        else:
            return True, "CV uploaded successfully, but no text could be extracted for auto-fill."
    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Extraction error: {e}")
        return True, "CV uploaded successfully, but an error occurred during auto-fill extraction."
